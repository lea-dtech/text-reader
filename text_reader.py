import sys
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QDialog, QApplication, QFileDialog
from PyQt5.uic import loadUi

# docx Table
import docx
from docx.shared import Inches

class MainWindow(QDialog):
    def __init__(self):
        super(MainWindow, self).__init__()
        loadUi("gui.ui",self)
        self.Browse.clicked.connect(self.browsefiles)

    def browsefiles(self):
        fname=QFileDialog.getOpenFileName(self, "open file", "c:/users",'TXT files(*.txt)')
        fpath=fname[0]
        if fname[0]!="":
            f=fpath
            self.filterline(f)
        else:
            print("Please Choose file..!")

    def filterline(self,f):
        f=open(f)
        doc= docx.Document()
        sNo=0
        data=[]
        for line in f:
            if ("not ok" or "NOT OK") in line:
                sNo+=1
                doc.add_paragraph(str(sNo)+". "+line)
                data.append(line)
        # Table codes for docx file
        sections=doc.sections
        for section in sections:
            section.top_margin=Inches(0.5)
            section.bottom_margin=Inches(1)
            section.left_margin=Inches(1)
            section.right_margin=Inches(1)

        doc.add_heading("Heading",0)
        #Table's colum heading
        table=doc.add_table(rows=1,cols=4)
        table.style='Table Grid'
        hdr_cells=table.rows[0].cells
        hdr_cells[0].text="S. NO"
        hdr_cells[1].text='Exceptation'
        hdr_cells[2].text='Obsjerved'
        hdr_cells[3].text='Remarks'

        s_no=0
        for line in data:
            s_no+=1
            row_cells=table.add_row().cells
            row_cells[0].text=str(s_no)
            row_cells[1].text=line
            row_cells[2].text=line
            row_cells[3].text="pass"



        f.close()
        doc.save("test1.docx")
        print("Operation Success..!")
        

app=QApplication(sys.argv)

Mainwindow=MainWindow()
widget=QtWidgets.QStackedWidget()
widget.addWidget(Mainwindow)
widget.setFixedWidth(400)
widget.setFixedHeight(300)
widget.show()
sys.exit(app.exec_())
