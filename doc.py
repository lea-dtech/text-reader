import docx
from docx.shared import Inches

doc= docx.Document()

sections=doc.sections
for section in sections:
    section.top_margin=Inches(0.5)
    section.bottom_margin=Inches(1)
    section.left_margin=Inches(1)
    section.right_margin=Inches(1)

doc.add_heading("Heading",0)

table=doc.add_table(rows=1,cols=4)
table.style='Table Grid'
hdr_cells=table.rows[0].cells
hdr_cells[0].text="S. NO"
hdr_cells[1].text='Exceptation'
hdr_cells[2].text='Obsjerved'
hdr_cells[3].text='Remarks'


doc.save("test.docx")